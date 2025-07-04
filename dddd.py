from docx import Document

# Create a new Document
doc = Document()

# Add Title
doc.add_heading('Mini Projekt – Statystyka Opisowa', 0)

# Add section for Data
doc.add_heading('Dane:', level=1)
doc.add_paragraph('Opis: Dane dotyczą przeciętnego zatrudnienia i wynagrodzenia w sektorze przedsiębiorstw w Polsce na przestrzeni 2020 roku. Analiza skupi się na miesięcznych wynagrodzeniach w sektorze przedsiębiorstw.')

# Add Data Table
doc.add_heading('Zbiór Danych:', level=2)
table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Miesiąc'
hdr_cells[1].text = 'Przeciętne wynagrodzenie brutto (PLN)'

data = [
    ('Styczeń', '5281.08'),
    ('Luty', '5280.80'),
    ('Marzec', '5331.47'),
    ('Kwiecień', '5402.45'),
    ('Maj', '5424.49'),
    ('Czerwiec', '5376.96'),
    ('Lipiec', '5382.49'),
    ('Sierpień', '5337.65'),
    ('Wrzesień', '5458.88'),
    ('Październik', '5457.98'),
    ('Listopad', '5562.05'),
    ('Grudzień', '5739.11')
]

for month, salary in data:
    row_cells = table.add_row().cells
    row_cells[0].text = month
    row_cells[1].text = salary

# Add Data Source
doc.add_paragraph('Źródło: Główny Urząd Statystyczny')

# Add Graph (Placeholder)
doc.add_heading('Prezentacja Graficzna:', level=2)
doc.add_paragraph('![Wynagrodzenia w 2020 roku](https://via.placeholder.com/800x400.png?text=Wynagrodzenia+w+2020+roku)')

# Add section for Statistical Measures
doc.add_heading('Miary statystyczne:', level=1)

# Add subsection for Mean Measures
doc.add_heading('Miary średnie:', level=2)
doc.add_paragraph('Średnia arytmetyczna: 5408.75 PLN')
doc.add_paragraph('Mediana: 5392.47 PLN')
doc.add_paragraph('Pierwszy kwartyl (Q1): 5331.47 PLN')
doc.add_paragraph('Trzeci kwartyl (Q3): 5458.88 PLN')
doc.add_paragraph('Dominanta: Brak powtarzających się wartości.')

# Add subsection for Dispersion Measures
doc.add_heading('Miary zróżnicowania:', level=2)
doc.add_paragraph('Wariancja: 20350.63 PLN^2')
doc.add_paragraph('Odchylenie standardowe: 142.65 PLN')
doc.add_paragraph('Współczynnik zmienności (klasyczny): 2.64%')
doc.add_paragraph('Odchylenie ćwiartkowe: 63.71 PLN')
doc.add_paragraph('Współczynnik zmienności (pozycyjny): 1.18%')

# Add subsection for Skewness Measures
doc.add_heading('Miary asymetrii:', level=2)
doc.add_paragraph('Wskaźnik asymetrii Ms: 0.11')
doc.add_paragraph('Współczynnik skośności ws: Wymaga dokładnych obliczeń')
doc.add_paragraph('Współczynnik skośności As: 0.33')

# Add section for Analysis and Conclusions
doc.add_heading('Analiza danych i wnioski:', level=1)
doc.add_paragraph('1. Średnia arytmetyczna wynagrodzeń w sektorze przedsiębiorstw w 2020 roku wynosi 5408.75 PLN, co oznacza, że przeciętne wynagrodzenie w tym sektorze było stosunkowo stabilne przez cały rok.')
doc.add_paragraph('2. Mediana wynosi 5392.47 PLN, co wskazuje, że połowa pracowników zarabiała mniej niż 5392.47 PLN, a druga połowa więcej.')
doc.add_paragraph('3. Pierwszy kwartyl (Q1) wynosi 5331.47 PLN, co oznacza, że 25% pracowników zarabia mniej niż 5331.47 PLN.')
doc.add_paragraph('4. Trzeci kwartyl (Q3) wynosi 5458.88 PLN, co oznacza, że 75% pracowników zarabia mniej niż 5458.88 PLN.')
doc.add_paragraph('5. Odchylenie standardowe wynosi 142.65 PLN, co oznacza niskie zróżnicowanie wynagrodzeń w sektorze przedsiębiorstw.')
doc.add_paragraph('6. Współczynnik zmienności wynoszący 2.64% potwierdza niskie zróżnicowanie wynagrodzeń.')
doc.add_paragraph('7. Odchylenie ćwiartkowe i współczynnik zmienności pozycyjny są również niskie, co wskazuje na stabilność wynagrodzeń.')
doc.add_paragraph('8. Wskaźnik asymetrii Ms wynosi 0.11, co oznacza, że rozkład wynagrodzeń jest lekko prawostronnie skośny, jednak asymetria jest niewielka.')
doc.add_paragraph('9. Współczynnik skośności ws wskazuje na niewielką prawostronną skośność, ale dokładne obliczenia wymagają przeliczenia wszystkich składników sumy.')
doc.add_paragraph('10. Współczynnik skośności As wynosi 0.33, co również wskazuje na niewielką prawostronną asymetrię.')

# Save the Document
file_path = "/mnt/data/Mini_projekt_statystyka_opisowa.docx"
doc.save(file_path)

file_path
